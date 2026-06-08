"""
comparacion_vs_competicion_docx.py
==================================
Genera el documento Word de comparación entre NUESTROS resultados bajo
reglas de competición y los RESULTADOS OFICIALES del IHTC 2024.

Entradas:
  1. resultados_competicion_benchmark/detalle_iXX.json
        → nuestros resultados bajo reglas (1 hilo, 600 s). Generados por
          ejecutar_competicion.py. Si no existen aún, la columna "nuestro
          coste" aparece como 'pendiente'.
  2. resultados_oficiales.json
        → costes oficiales por instancia (best_known). Si están a null,
          la columna oficial aparece como '[rellenar]' y el gap como '—'.

Gap (cuando hay ambos datos):  gap% = (nuestro − oficial) / oficial × 100

El documento NO inventa ningún número: lo que falta se marca como pendiente.

Salida: resultados_competicion_benchmark/comparacion_vs_competicion.docx

Uso:
  python comparacion_vs_competicion_docx.py
  python comparacion_vs_competicion_docx.py --limite 600
"""

import argparse
import glob
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent.resolve()
AZUL = RGBColor(0x2E, 0x57, 0x9C)
ROJO = RGBColor(0xC0, 0x00, 0x00)
VERDE = RGBColor(0x1F, 0x7A, 0x1F)
GRIS_BG = "F2F2F2"


def _sombrear(celda, hexcolor):
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _set(celda, texto, bold=False, size=9, align="center", color=None, bg=None):
    celda.text = ""
    p = celda.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r = p.add_run(str(texto))
    r.bold = bold
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    if bg is not None:
        _sombrear(celda, bg)


def _cab(tabla, headers):
    fila = tabla.rows[0].cells
    for i, h in enumerate(headers):
        _set(fila[i], h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF), bg="2E579C")


def cargar_nuestros(carpeta):
    res = {}
    for f in sorted(glob.glob(str(carpeta / "detalle_i*.json"))):
        d = json.load(open(f, encoding="utf-8"))
        res[d.get("instancia")] = d
    return res


def main():
    for _s in (sys.stdout, sys.stderr):
        try:
            _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

    ap = argparse.ArgumentParser()
    ap.add_argument("--carpeta", default="resultados_competicion_benchmark")
    ap.add_argument("--oficiales", default="resultados_oficiales.json")
    ap.add_argument("--limite", type=float, default=600.0)
    ap.add_argument("--etiqueta", default="reglas de competición (4 hilos, tope 600 s)",
                    help="Descripción de las condiciones de NUESTRO run para el título.")
    ap.add_argument("--preview", action="store_true",
                    help="Marca el documento como comparación PRELIMINAR (condiciones "
                         "ampliadas, no homologable). Añade un aviso destacado.")
    ap.add_argument("--salida", default="comparacion_vs_competicion.docx")
    args = ap.parse_args()

    carpeta = BASE_DIR / args.carpeta
    nuestros = cargar_nuestros(carpeta) if carpeta.exists() else {}

    of_path = BASE_DIR / args.oficiales
    oficiales = {}
    if of_path.exists():
        oficiales = json.load(open(of_path, encoding="utf-8")).get("instancias", {})

    instancias = [f"i{n:02d}" for n in range(1, 31)]

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    titulo = ("Comparación PRELIMINAR con los resultados oficiales del IHTC 2024"
              if args.preview else
              "Comparación con los resultados oficiales del IHTC 2024")
    t = doc.add_heading(titulo, level=0)
    for r in t.runs:
        r.font.color.rgb = AZUL
    p = doc.add_paragraph()
    p.add_run(f"Resultados de la propuesta en {args.etiqueta} frente al mejor resultado "
              "por instancia publicado por los participantes "
              "(reglas oficiales: 10 min y máximo 4 hilos).").italic = True

    if args.preview:
        av = doc.add_paragraph()
        rr = av.add_run(
            "⚠ COMPARACIÓN PRELIMINAR / OPTIMISTA: nuestros resultados se obtuvieron en "
            "CONDICIONES AMPLIADAS (8 hilos y sin límite estricto de tiempo), superiores a "
            "las de competición. Los gaps mostrados son por tanto una COTA INFERIOR del "
            "gap real; bajo reglas estrictas (4 hilos, 600 s) el gap será mayor. Para la "
            "comparación homologable, ejecutar «ejecutar_competicion.py» y regenerar.")
        rr.bold = True
        rr.font.color.rgb = ROJO

    # Estado de los datos
    hay_nuestros = len(nuestros) > 0
    hay_oficiales = any(v.get("best_known") is not None for v in oficiales.values())

    doc.add_heading("1. Estado de los datos", level=1)
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"Resultados propios ({args.etiqueta}): ").bold = True
    if hay_nuestros:
        n_ok = sum(1 for d in nuestros.values() if d.get("estado") == "OK")
        p.add_run(f"disponibles ({n_ok}/{len(nuestros)} resueltas).")
    else:
        p.add_run("PENDIENTES. Ejecuta «python ejecutar_competicion.py» para generarlos.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Resultados oficiales IHTC 2024: ").bold = True
    if hay_oficiales:
        n = sum(1 for v in oficiales.values() if v.get("best_known") is not None)
        p.add_run(f"cargados para {n}/30 instancias desde resultados_oficiales.json.")
    else:
        p.add_run("PENDIENTES. Rellena «resultados_oficiales.json» con los costes oficiales.")

    # Tabla comparativa
    doc.add_heading("2. Comparación por instancia", level=1)
    headers = ["Inst", "Nuestro coste\n(reglas)", "Violac.", "t (s)",
               "Resultado\noficial", "Gap %", "Fuente oficial"]
    tabla = doc.add_table(rows=1, cols=len(headers))
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cab(tabla, headers)

    gaps = []
    for idx, inst in enumerate(instancias):
        bg = GRIS_BG if idx % 2 == 1 else None
        d = nuestros.get(inst)
        of = oficiales.get(inst, {})
        bks = of.get("best_known")
        fuente = of.get("fuente", "") or ""

        # Nuestro coste
        if d and d.get("estado") == "OK":
            nuestro = d.get("total_cost")
            viol = d.get("total_violations", 0)
            t_s = d.get("tiempos", {}).get("total")
            nuestro_str = f"{nuestro:,}".replace(",", ".")
            t_str = f"{t_s:.0f}" if t_s is not None else "–"
        elif d:
            nuestro = None
            viol = d.get("estado", "?")
            nuestro_str = "–"
            t_str = "–"
        else:
            nuestro = None
            viol = "–"
            nuestro_str = "pendiente"
            t_str = "–"

        # Oficial + gap
        if bks is not None:
            of_str = f"{bks:,}".replace(",", ".")
        else:
            of_str = "[rellenar]"

        if nuestro is not None and bks:
            gap = (nuestro - bks) / bks * 100
            gaps.append(gap)
            gap_str = f"{gap:+.1f}%"
            gap_color = VERDE if gap <= 0 else (ROJO if gap > 25 else None)
        else:
            gap_str = "—"
            gap_color = None

        vals = [inst, nuestro_str, str(viol), t_str, of_str, gap_str, fuente]
        celdas = tabla.add_row().cells
        for i, v in enumerate(vals):
            color = gap_color if i == 5 else None
            _set(celdas[i], v, size=9, bold=(i == 1), color=color, bg=bg,
                 align="left" if i == 6 else "center")

    # Resumen
    doc.add_heading("3. Resumen", level=1)
    if gaps:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("Gap medio vs oficial: ").bold = True
        p.add_run(f"{sum(gaps)/len(gaps):+.2f}%  (sobre {len(gaps)} instancias con dato oficial)")
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("Gap mínimo / máximo: ").bold = True
        p.add_run(f"{min(gaps):+.1f}% / {max(gaps):+.1f}%")
        n_mejor = sum(1 for g in gaps if g <= 0)
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("Instancias que igualan o mejoran el oficial: ").bold = True
        p.add_run(f"{n_mejor}/{len(gaps)}")
    else:
        p = doc.add_paragraph()
        p.add_run("Resumen pendiente: ").bold = True
        p.add_run("se completará automáticamente cuando estén disponibles nuestros "
                  "resultados bajo reglas (ejecutar_competicion.py) y los resultados "
                  "oficiales (resultados_oficiales.json).").italic = True

    p = doc.add_paragraph()
    p.add_run("Nota: ").bold = True
    p.add_run("este documento no contiene ningún valor inventado. Las celdas "
              "«pendiente» y «[rellenar]» se sustituirán automáticamente al regenerar "
              "el documento con los datos reales.").italic = True

    carpeta.mkdir(parents=True, exist_ok=True)
    ruta = carpeta / args.salida
    doc.save(str(ruta))
    print(f"Documento generado: {ruta}")
    print(f"Nuestros resultados: {'SÍ' if hay_nuestros else 'pendientes'} | "
          f"Oficiales: {'SÍ' if hay_oficiales else 'pendientes'}")


if __name__ == "__main__":
    main()
