"""
ablacion_tabla_docx.py
======================
Genera la tabla de ablación (Word + impresión) comparando dos variantes del
pipeline sobre las MISMAS instancias:

  · CON readmisión   → resultados_rapido_benchmark/        (ejecutar_rapido.py)
  · SIN readmisión   → resultados_ablacion_benchmark/      (ejecutar_ablacion_sin_readmision.py)

Para cada instancia (presente en ambos) muestra:
  coste sin readmisión · coste con readmisión · mejora absoluta · mejora %
  · nº de opcionales podados en cada variante.

Solo compara instancias resueltas (OK) en AMBAS variantes; el resto se omite
del cálculo de medias (se indica aparte).

Salida: resultados_ablacion_benchmark/ablacion_readmision.docx

Uso:
  python ablacion_tabla_docx.py
  python ablacion_tabla_docx.py --con resultados_rapido_benchmark --sin resultados_ablacion_benchmark
"""

import argparse
import glob
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent.resolve()
AZUL = RGBColor(0x2E, 0x57, 0x9C)
VERDE = RGBColor(0x1F, 0x7A, 0x1F)
ROJO = RGBColor(0xC0, 0x00, 0x00)
GRIS_BG = "F2F2F2"


def cargar(carpeta):
    res = {}
    for f in sorted(glob.glob(str(carpeta / "detalle_i*.json"))):
        d = json.load(open(f, encoding="utf-8"))
        res[d.get("instancia")] = d
    return res


def _shd(c, hexcolor):
    tcPr = c._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear"); s.set(qn("w:fill"), hexcolor)
    tcPr.append(s)


def _set(c, txt, bold=False, size=9, align="center", color=None, bg=None):
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r = p.add_run(str(txt)); r.bold = bold; r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    if bg is not None: _shd(c, bg)


def _cab(tabla, headers):
    for i, h in enumerate(headers):
        _set(tabla.rows[0].cells[i], h, bold=True, size=9,
             color=RGBColor(0xFF, 0xFF, 0xFF), bg="2E579C")


def main():
    for _s in (sys.stdout, sys.stderr):
        try: _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception: pass

    ap = argparse.ArgumentParser()
    ap.add_argument("--con", default="resultados_rapido_benchmark")
    ap.add_argument("--sin", default="resultados_ablacion_benchmark")
    args = ap.parse_args()

    con = cargar(BASE_DIR / args.con)
    sin = cargar(BASE_DIR / args.sin)

    comunes = sorted(set(con) & set(sin))
    if not comunes:
        print("[ERROR] No hay instancias comunes entre las dos carpetas.")
        print("  Genera primero la variante sin readmisión con "
              "ejecutar_ablacion_sin_readmision.py")
        return

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    t = doc.add_heading("Estudio de ablación: aportación de la readmisión", level=0)
    for r in t.runs: r.font.color.rgb = AZUL
    p = doc.add_paragraph()
    p.add_run("Comparación del pipeline CON y SIN la fase de readmisión de pacientes en "
              "la poda, manteniendo idénticos el resto de parámetros (tiempos, semilla, "
              "hilos). El delta de coste se atribuye exclusivamente a la readmisión.").italic = True

    headers = ["Inst", "Coste SIN\nreadmisión", "Podados\nSIN", "Coste CON\nreadmisión",
               "Podados\nCON", "Mejora", "Mejora %"]
    tabla = doc.add_table(rows=1, cols=len(headers))
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cab(tabla, headers)

    mejoras_pct = []
    pod_sin_tot = pod_con_tot = 0
    filas_validas = 0
    idx = 0
    for inst in comunes:
        ds, dc = sin[inst], con[inst]
        if ds.get("estado") != "OK" or dc.get("estado") != "OK":
            continue
        cs = ds.get("total_cost"); cc = dc.get("total_cost")
        if cs is None or cc is None:
            continue
        ps = ds.get("metricas", {}).get("pacientes_eliminados_poda")
        pc = dc.get("metricas", {}).get("pacientes_eliminados_poda")
        ps = 0 if ps is None else ps
        pc = 0 if pc is None else pc
        # Solo interesa donde hubo poda en alguna variante
        if ps == 0 and pc == 0:
            continue

        mejora = cs - cc
        mejora_pct = mejora / cs * 100 if cs else 0
        mejoras_pct.append(mejora_pct)
        pod_sin_tot += ps; pod_con_tot += pc
        filas_validas += 1

        bg = GRIS_BG if idx % 2 == 1 else None
        idx += 1
        color = VERDE if mejora > 0 else (ROJO if mejora < 0 else None)
        vals = [inst, f"{cs:,}".replace(",", "."), ps, f"{cc:,}".replace(",", "."), pc,
                f"{mejora:+,}".replace(",", "."), f"{mejora_pct:+.1f}%"]
        celdas = tabla.add_row().cells
        for i, v in enumerate(vals):
            _set(celdas[i], v, size=9, bold=(i in (5, 6)),
                 color=color if i in (5, 6) else None, bg=bg)

    doc.add_heading("Resumen", level=1)
    if mejoras_pct:
        def b(label, val):
            pp = doc.add_paragraph(style="List Bullet")
            pp.add_run(label + ": ").bold = True
            pp.add_run(val)
        b("Instancias con poda comparadas", str(filas_validas))
        b("Mejora media de coste por la readmisión", f"{sum(mejoras_pct)/len(mejoras_pct):+.1f}%")
        b("Mejor caso", f"{max(mejoras_pct):+.1f}%")
        b("Total opcionales podados", f"SIN readmisión: {pod_sin_tot}  →  CON: {pod_con_tot} "
          f"(recuperados: {pod_sin_tot - pod_con_tot})")
        n_mejora = sum(1 for m in mejoras_pct if m > 0)
        b("Instancias donde la readmisión mejora", f"{n_mejora}/{filas_validas}")
    else:
        doc.add_paragraph("Sin datos comparables todavía. Ejecuta la variante sin "
                          "readmisión sobre las instancias que pasan por poda.")

    carpeta_out = BASE_DIR / args.sin
    carpeta_out.mkdir(parents=True, exist_ok=True)
    ruta = carpeta_out / "ablacion_readmision.docx"
    doc.save(str(ruta))

    print(f"Documento generado: {ruta}")
    print(f"Instancias comparadas (con poda): {filas_validas}")
    if mejoras_pct:
        print(f"Mejora media por readmisión: {sum(mejoras_pct)/len(mejoras_pct):+.1f}%")
        print(f"Opcionales podados: SIN {pod_sin_tot} -> CON {pod_con_tot}")


if __name__ == "__main__":
    main()
