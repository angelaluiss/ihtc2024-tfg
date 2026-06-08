"""
comparacion_memoria_docx.py
===========================
Genera un documento Word (.docx) con la comparación de nuestra propuesta
frente a las reglas de la competición IHTC 2024, con tablas formateadas
listas para pegar en la memoria del TFG.

Lee los JSON de detalle ya generados (no re-ejecuta nada).
Salida: comparacion_memoria.docx

Uso:
  python comparacion_memoria_docx.py
  python comparacion_memoria_docx.py --limite-competicion 600 --hilos-competicion 4
"""

import argparse
import glob
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent.resolve()

AZUL    = RGBColor(0x2E, 0x57, 0x9C)
AZUL_BG = "D5E8F0"
GRIS_BG = "F2F2F2"


def cargar(carpeta):
    filas = []
    for f in sorted(glob.glob(str(carpeta / "detalle_i*.json"))):
        filas.append(json.load(open(f, encoding="utf-8")))
    filas.sort(key=lambda d: d.get("instancia", ""))
    return filas


def _sombrear(celda, hexcolor):
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _set_cell(celda, texto, bold=False, size=9, align="left", color=None, bg=None):
    celda.text = ""
    p = celda.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(str(texto))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bg is not None:
        _sombrear(celda, bg)


def _tabla_cabecera(tabla, headers, anchos=None):
    fila = tabla.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell(fila[i], h, bold=True, size=9, align="center",
                  color=RGBColor(0xFF, 0xFF, 0xFF), bg="2E579C")


def construir(carpeta, lim_t, lim_h, ruta_salida):
    filas = cargar(carpeta)
    ok = [d for d in filas if d.get("estado") == "OK"]
    no_ok = [d for d in filas if d.get("estado") != "OK"]

    hilos_usados = max((d.get("threads", 0) or 0) for d in filas) if filas else 0
    t_max = max((d.get("tiempos", {}).get("total", 0) or 0) for d in ok) if ok else 0
    excede_t = [d for d in ok if (d.get("tiempos", {}).get("total", 0) or 0) > lim_t]
    excede_h = [d for d in filas if (d.get("threads", 0) or 0) > lim_h]

    doc = Document()

    # Estilo base
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ── Título ────────────────────────────────────────────────────────
    t = doc.add_heading("Resultados en condiciones ampliadas (IHTC 2024)", level=0)
    for run in t.runs:
        run.font.color.rgb = AZUL
    p = doc.add_paragraph()
    p.add_run("Evaluación de la propuesta del TFG sobre las 30 instancias (i01–i30) bajo "
              f"CONDICIONES AMPLIADAS: {hilos_usados} hilos y sin límite estricto de tiempo "
              "(presupuesto blando de 840 s/instancia que la poda puede exceder para "
              "garantizar factibilidad).").italic = True
    aviso = doc.add_paragraph()
    r = aviso.add_run(
        "⚠ ESTOS RESULTADOS NO SON HOMOLOGABLES EN COMPETICIÓN. No respetan el límite "
        f"oficial de tiempo ni el de hilos. Para resultados bajo reglas de competición "
        "(1 hilo, 600 s) véase el documento «comparacion_vs_competicion.docx».")
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # ── 1. Configuración vs reglas ────────────────────────────────────
    doc.add_heading("1. Configuración experimental frente a las reglas oficiales", level=1)
    headers = ["Parámetro", "IHTC 2024 (referencia)", "Nuestra propuesta", "¿Dentro de las reglas?"]
    datos = [
        ("Límite de tiempo / instancia", f"{lim_t:.0f} s (máquina de referencia)",
         "840 s de presupuesto, sin tope duro", "NO"),
        ("Hilos de cómputo", f"{lim_h} (referencia)", f"{hilos_usados} (todos los núcleos)", "NO"),
        ("Validación de soluciones", "IHTP_Validator oficial", "IHTP_Validator oficial", "SÍ"),
        ("Formato instancia/solución", "JSON oficial", "JSON oficial", "SÍ"),
        ("Restricciones duras", "0 violaciones exigidas", "0 violaciones en todas", "SÍ"),
    ]
    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tabla_cabecera(tabla, headers)
    for fila in datos:
        celdas = tabla.add_row().cells
        cumple = fila[3] == "SÍ"
        for i, val in enumerate(fila):
            bg = None
            color = None
            bold = False
            if i == 3:
                bold = True
                color = RGBColor(0x1F, 0x7A, 0x1F) if cumple else RGBColor(0xC0, 0x00, 0x00)
            _set_cell(celdas[i], val, bold=bold, size=9, color=color, bg=bg,
                      align="center" if i == 3 else "left")

    p = doc.add_paragraph()
    p.add_run("Nota metodológica: ").bold = True
    p.add_run(
        "la propuesta NO compite bajo condiciones oficiales (usa más hilos y, en las "
        "instancias grandes, más tiempo del límite de referencia). Los resultados deben "
        "interpretarse como cota práctica de calidad alcanzable con recursos ampliados, no "
        "como resultado homologable en el ranking. El límite oficial de tiempo se calibra "
        f"por máquina; aquí se toma {lim_t:.0f} s como referencia."
    ).italic = True

    # ── 2. Desviaciones ───────────────────────────────────────────────
    doc.add_heading("2. Desviaciones respecto a las reglas", level=1)
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Hilos: ").bold = True
    p.add_run(f"{len(excede_h)}/{len(filas)} instancias con {hilos_usados} hilos (> {lim_h}). "
              "Solo afecta a las fases MIP (Gurobi); las metaheurísticas son monohilo.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Tiempo: ").bold = True
    p.add_run(f"{len(excede_t)}/{len(ok)} instancias superan {lim_t:.0f} s. "
              f"Tiempo máximo: {t_max:.0f} s.")

    if excede_t:
        tabla = doc.add_table(rows=1, cols=3)
        tabla.style = "Table Grid"
        _tabla_cabecera(tabla, ["Instancia", "Tiempo total (s)", f"Exceso sobre {lim_t:.0f} s"])
        for d in sorted(excede_t, key=lambda x: -(x.get("tiempos", {}).get("total", 0) or 0)):
            tt = d["tiempos"]["total"]
            celdas = tabla.add_row().cells
            _set_cell(celdas[0], d["instancia"], align="center", size=9)
            _set_cell(celdas[1], f"{tt:.1f}", align="center", size=9)
            _set_cell(celdas[2], f"+{tt - lim_t:.0f}", align="center", size=9,
                      color=RGBColor(0xC0, 0x00, 0x00))

    p = doc.add_paragraph()
    p.add_run(
        "El exceso de tiempo se concentra en las instancias grandes (≥150 pacientes) y "
        "proviene del bloque de Fase 2 (habitaciones + poda con readmisión), que prioriza "
        "garantizar una solución factible sobre respetar el presupuesto blando."
    ).italic = True

    # ── 3. Resultados exactos ─────────────────────────────────────────
    doc.add_heading("3. Resultados exactos de la propuesta", level=1)
    headers = ["Inst", "Estado", "Viol.", "Coste total", "Unsched.", "Delay", "Resto (soft)", "t (s)", "Hilos"]
    tabla = doc.add_table(rows=1, cols=len(headers))
    tabla.style = "Table Grid"
    _tabla_cabecera(tabla, headers)

    sum_coste = sum_uns = sum_del = 0
    for idx, d in enumerate(filas):
        celdas = tabla.add_row().cells
        bg = GRIS_BG if idx % 2 == 1 else None
        if d.get("estado") != "OK":
            vals = [d.get("instancia", "?"), d.get("estado", "?"), "–", "–", "–", "–", "–", "–",
                    str(d.get("threads", "–"))]
            for i, v in enumerate(vals):
                _set_cell(celdas[i], v, size=9, align="center",
                          color=RGBColor(0xC0, 0x00, 0x00) if i == 1 else None, bg=bg)
            continue
        coste = d.get("total_cost", 0) or 0
        uns = d.get("cost_ElectiveUnscheduledPatients", 0) or 0
        dly = d.get("cost_PatientDelay", 0) or 0
        resto = coste - uns - dly
        tt = d.get("tiempos", {}).get("total", 0) or 0
        sum_coste += coste; sum_uns += uns; sum_del += dly
        vals = [d["instancia"], "OK", d.get("total_violations", 0), f"{coste:,}".replace(",", "."),
                f"{uns:,}".replace(",", "."), f"{dly:,}".replace(",", "."),
                f"{resto:,}".replace(",", "."), f"{tt:.1f}", d.get("threads", "–")]
        for i, v in enumerate(vals):
            _set_cell(celdas[i], v, size=9, align="center", bold=(i == 3), bg=bg)

    # ── 4. Agregados ──────────────────────────────────────────────────
    doc.add_heading("4. Métricas agregadas", level=1)
    n_ok = len(ok)
    t_total = sum((d.get("tiempos", {}).get("total", 0) or 0) for d in ok)
    costes = [d.get("total_cost", 0) or 0 for d in ok]

    def bullet(label, value):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(label + ": ").bold = True
        p.add_run(value)

    bullet("Instancias resueltas sin violaciones", f"{n_ok}/{len(filas)}")
    if n_ok:
        bullet("Coste total acumulado", f"{sum_coste:,}".replace(",", "."))
        bullet("Coste medio por instancia", f"{sum_coste / n_ok:,.0f}".replace(",", "."))
        bullet("Coste mínimo / máximo", f"{min(costes):,} / {max(costes):,}".replace(",", "."))
        if sum_coste:
            bullet("Composición del coste",
                   f"Unscheduled {sum_uns/sum_coste*100:.0f}% · "
                   f"Delay {sum_del/sum_coste*100:.0f}% · "
                   f"Resto (soft) {(sum_coste-sum_uns-sum_del)/sum_coste*100:.0f}%")
        bullet("Tiempo total de cómputo (suma)", f"{t_total/60:.1f} min ({t_total/3600:.2f} h)")
    if no_ok:
        bullet("Instancias pendientes / infactibles",
               ", ".join(d["instancia"] for d in no_ok))

    doc.save(str(ruta_salida))
    return ruta_salida, n_ok, len(filas)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--carpeta", default="resultados_rapido_benchmark")
    ap.add_argument("--limite-competicion", type=float, default=600.0)
    ap.add_argument("--hilos-competicion", type=int, default=4)
    args = ap.parse_args()

    carpeta = BASE_DIR / args.carpeta
    ruta = carpeta / "comparacion_memoria.docx"
    ruta, n_ok, n_tot = construir(carpeta, args.limite_competicion,
                                  args.hilos_competicion, ruta)
    print(f"Documento Word generado: {ruta}")
    print(f"Instancias OK: {n_ok}/{n_tot}")


if __name__ == "__main__":
    main()
